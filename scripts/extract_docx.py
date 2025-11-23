from docx import Document
import sys

# Read the DOCX file
doc = Document(r'C:\Users\Burak\Desktop\Stratejiler5_ornek_sayfalar.docx')

# Extract all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

# Also extract text from tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text)

# Write to file with UTF-8 encoding
with open('exam_content.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(full_text))

print(f"Extracted {len(full_text)} lines from DOCX")
print("First 10 lines:")
for i, line in enumerate(full_text[:10]):
    print(f"{i+1}. {line[:100]}")
